import io
import re

from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document
import pytesseract

from ocr.ocr_pretraitement import preprocess_image
from ocr.all_type import parse_document_with_ai


INVOICE_KEYWORDS = [
    "facture",
    "invoice",
    "date",
    "total",
    "tva",
    "vat",
    "montant",
    "amount",
    "prix",
    "subtotal",
    "sous-total",
    "référence",
    "reference",
    "client",
    "commande",
    "order",
]


def _normalize_text(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text).strip().lower()


def _extract_amounts(text: str) -> list[str]:
    if not text:
        return []
    return re.findall(r"\d+[.,]\d{2}", text)


def _count_keyword_hits(text: str) -> int:
    lowered = _normalize_text(text)
    return sum(1 for kw in INVOICE_KEYWORDS if kw in lowered)


def _count_weird_chars(text: str) -> int:
    if not text:
        return 0
    weird_set = set("�□{}[]|<>~")
    return sum(1 for c in text if c in weird_set)


def _line_count(text: str) -> int:
    if not text:
        return 0
    return len([line for line in text.splitlines() if line.strip()])


def _alpha_count(text: str) -> int:
    if not text:
        return 0
    return sum(1 for c in text if c.isalpha())


def ocr_quality_score(text: str) -> float:
    """
    Score heuristique de qualité OCR entre 0 et 1.
    Plus le score est élevé, plus le texte semble exploitable.
    """
    if not text or not text.strip():
        return 0.0

    cleaned = text.strip()
    score = 0.0

    # 1. Longueur
    length = len(cleaned)
    if length >= 50:
        score += 0.15
    if length >= 120:
        score += 0.10
    if length >= 300:
        score += 0.10

    # 2. Nombre de lignes
    lines = _line_count(cleaned)
    if lines >= 3:
        score += 0.10
    if lines >= 8:
        score += 0.08

    # 3. Présence de lettres
    alpha = _alpha_count(cleaned)
    if alpha >= 20:
        score += 0.10
    if alpha >= 80:
        score += 0.07

    # 4. Mots-clés facture
    keyword_hits = _count_keyword_hits(cleaned)
    if keyword_hits >= 2:
        score += 0.12
    if keyword_hits >= 4:
        score += 0.08
    if keyword_hits >= 6:
        score += 0.05

    # 5. Montants
    amounts = _extract_amounts(cleaned)
    if len(amounts) >= 1:
        score += 0.10
    if len(amounts) >= 3:
        score += 0.05

    # 6. Pénalités caractères incohérents
    weird_chars = _count_weird_chars(cleaned)
    if weird_chars >= 5:
        score -= 0.10
    if weird_chars >= 10:
        score -= 0.10

    # 7. Pénalité si trop peu structuré
    if lines < 2:
        score -= 0.10

    return max(0.0, min(score, 1.0))


def is_ocr_result_bad(text: str, threshold: float = 0.65) -> bool:
    return ocr_quality_score(text) < threshold


def compare_text_candidates(local_text: str, ai_text: str) -> str:
    """
    Compare deux résultats et retourne le meilleur.
    On utilise un score heuristique + un léger bonus au texte plus complet.
    """
    local_score = ocr_quality_score(local_text)
    ai_score = ocr_quality_score(ai_text)

    local_len = len(local_text.strip()) if local_text else 0
    ai_len = len(ai_text.strip()) if ai_text else 0

    # Bonus léger si le texte est plus riche/long
    local_final = local_score + min(local_len / 4000, 0.10)
    ai_final = ai_score + min(ai_len / 4000, 0.10)

    # Si l'IA est clairement meilleure, on la prend
    if ai_final > local_final + 0.05:
        return ai_text

    # Si les scores sont proches, on préfère le texte le plus complet
    if ai_len > local_len * 1.2:
        return ai_text

    return local_text


def extract_text_from_image(
    image: Image.Image,
    filename: str = "image",
    lang: str = "fra",
    force_ai_fallback: bool = False,
) -> str:
    """
    Pipeline pro :
    1. OCR local
    2. si mauvais -> appel GPT
    3. comparaison local vs GPT
    4. retour du meilleur
    """
    processed = preprocess_image(image)

    local_text = pytesseract.image_to_string(
        processed,
        lang=lang,
        config="--oem 3 --psm 6"
    ).strip()

    local_score = ocr_quality_score(local_text)

    print("LOCAL SCORE:", local_score)
    print("LOCAL TEXT:", local_text[:500])

    #  Si le texte est bon → on retourne direct (pas de coût GPT)
    if not force_ai_fallback and ocr_quality_score(local_text) >= 0.65:
        print("OCR local suffisant, pas de fallback GPT")
        return local_text

    print("Fallback GPT déclenché")

    #  Appel GPT
    ai_text = parse_document_with_ai(
        image=image,
        filename=filename,
    ).strip()

    print("AI TEXT:", ai_text[:500])

    # sécurité
    if not ai_text:
        print("GPT vide → fallback OCR local")
        return local_text

    #  Comparaison intelligente
    best_text = compare_text_candidates(local_text, ai_text)

    print("BEST SOURCE:", "GPT" if best_text == ai_text else "OCR")

    return best_text

def extract_text_from_pdf_bytes(
    contents: bytes,
    filename: str = "document.pdf",
    lang: str = "fra",
    force_ai_fallback: bool = False,
) -> str:
    pages = convert_from_bytes(contents, dpi=300)
    page_texts = []

    for page_number, page in enumerate(pages, start=1):
        best_text = extract_text_from_image(
            image=page,
            filename=f"{filename}_page_{page_number}",
            lang=lang,
            force_ai_fallback=force_ai_fallback,
        )
        page_texts.append(f"--- Page {page_number} ---\n{best_text}")

    return "\n\n".join(page_texts)


def extract_text_from_docx_bytes(contents: bytes) -> str:
    """
    Lecture native DOCX.
    """
    doc = Document(io.BytesIO(contents))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"--- Tableau {table_index} ---")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts).strip()